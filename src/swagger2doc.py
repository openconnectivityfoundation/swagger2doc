#############################
#
#    copyright 2016 Open Interconnect Consortium, Inc. All rights reserved.
#    Redistribution and use in source and binary forms, with or without modification,
#    are permitted provided that the following conditions are met:
#    1.  Redistributions of source code must retain the above copyright notice,
#        this list of conditions and the following disclaimer.
#    2.  Redistributions in binary form must reproduce the above copyright notice,
#        this list of conditions and the following disclaimer in the documentation and/or other materials provided
#        with the distribution.
#
#    THIS SOFTWARE IS PROVIDED BY THE OPEN INTERCONNECT CONSORTIUM, INC. "AS IS" AND ANY EXPRESS OR IMPLIED WARRANTIES,
#    INCLUDING, BUT NOT LIMITED TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE OR
#    WARRANTIES OF NON-INFRINGEMENT, ARE DISCLAIMED. IN NO EVENT SHALL THE OPEN INTERCONNECT CONSORTIUM, INC. OR
#    CONTRIBUTORS BE LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES
#    (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS;
#    OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
#    OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE,
#    EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
#############################


import time
import os
import json
import random
import sys
import argparse
import traceback
from datetime import datetime
from time import gmtime, strftime
#import jsonref

if sys.version_info < (3, 5):
    raise Exception("ERROR: Python 3.5 or more is required, you are currently running Python %d.%d!" %
                    (sys.version_info[0], sys.version_info[1]))

try:
    from docx import Document
except:
    print("missing swagger_parser:")
    print ("Trying to Install required module: python-docx (docx)")
    os.system('c:\python35\python.exe -m pip install python-docx')
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TABLESTYLE='TABLE-A'
TABLESTYLE='LightShading-Accent1'


# see https://github.com/python-openxml/python-docx/issues/359

def MarkIndexEntry(entry,paragraph):
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    #instrText.text = ' XE "%s" '%(entry)
    #r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def Figure(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def Table(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Table \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def load_json_schema(filename, my_dir):
    """
    load the JSON schema file
    :param filename: filename (with extension)
    :param my_dir: path to the file
    :return: json_dict
    """
    full_path = os.path.join(my_dir, filename)
    if os.path.isfile(full_path) is False:
        print ("json file does not exist:", full_path)

    linestring = open(full_path, 'r').read()
    json_dict = json.loads(linestring)

    return json_dict


def get_dir_list(dir, ext=None):
    """
    get all files (none recursive) in the specified dir
    :param dir: path to the directory
    :param ext: filter on extension
    :return: list of files (only base_name)
    """
    only_files = [f for f in listdir(dir) if isfile(join(dir, f))]
    # remove .bak files
    new_list = [x for x in only_files if not x.endswith(".bak")]
    if ext is not None:
        cur_list = new_list
        new_list = [x for x in cur_list if x.endswith(ext)]
    return new_list


def find_key(rec_dict, target, depth=0):
    """
    find key "target" in recursive dict
    :param rec_dict: dict to search in, json schema dict, so it is combination of dict and arrays
    :param target: target key to search for
    :param depth: depth of the search (recursion)
    :return:
    """
    try:
        if isinstance(rec_dict, dict):
            for key, value in rec_dict.items():
                if key == target:
                    return rec_dict[key]
            for key, value in rec_dict.items():
                r = find_key(value, target, depth+1)
                if r is not None:
                        return r
        #else:
        #    print ("no dict:", rec_dict)
    except:
        traceback.print_exc()


def find_key_link(rec_dict, target, depth=0):
    """
    find the first key recursively
    also traverse lists (arrays, oneOf,..) but only returns the first occurance
    :param rec_dict: dict to search in, json schema dict, so it is combination of dict and arrays
    :param target: target key to search for
    :param depth: depth of the search (recursion)
    :return:
    """
    if isinstance(rec_dict, dict):
        # direct key
        for key, value in rec_dict.items():
            if key == target:
                return rec_dict[key]
        # key is in array
        rvalues = []
        found = False
        for key, value in rec_dict.items():
            if key in ["oneOf", "allOf", "anyOf"]:
                for val in value:
                    if val == target:
                        return val
                    if isinstance(val, dict):
                        r = find_key_link(val, target, depth+1)
                        if r is not None:
                            found = True
                            # TODO: this should return an array, now it only returns the last found item
                            rvalues = r
        if found:
            return rvalues
        # key is an dict
        for key, value in rec_dict.items():
            r = find_key_link(value, target, depth+1)
            if r is not None:
                return r #[list(r.items())]


class CreateWordDoc(object):
    def __init__(self, swagger_parser, docx_name_in=None, docx_name_out=None, resource_name=None):
        """
        initialize the class


        """
        # input arguments
        self.docx_name_in = docx_name_in
        self.docx_name_out = docx_name_out
        self.resource_name = resource_name
        self.swagger_parser = swagger_parser
        self.annex_switch = False
        self.composite_switch = False
        self.sensor_switch = False
        self.schema_switch = False
        self.schemaWT_switch = False
        self.derived_name = None
        self.title = None

        try:
            schema_string = open(args.swagger, 'r').read()
        except:
            print ("CreateWordDoc *** ERROR : could not open:", args.swagger)
            return
        
        try:
            json_dict = json.loads(schema_string)
            self.json_parse_tree = json_dict
        except:
            print ("CreateWordDoc *** ERROR : error in JSON:", args.swagger)
            traceback.print_exc()
            

    def swag_sanitize_description(self, description):
        """
        removes line breaks, quotes  etc
        :param description: input string
        :return: text string
        """
        text = description.replace("\n", "@cr").replace("'", "<COMMA>").replace('"', "<COMMA>")
        return text

    def swag_unsanitize_description(self, description):
        """
        removes line breaks, quotes  etc
        :param description: input string
        :return: text string
        """
        text = description
        if description is not None:
            text = description.replace("@cr", "\n").replace("<COMMA>", "'")
        return text

    def parse_schema_requires(self, input_string_schema):
        """
        find the required property list

        :param input_string_schema: json schema as string
        :return:
        """
        ignore_list = ['required', '[', ']', ',', ': [']
        lines_schema = input_string_schema.splitlines()
        length = len(lines_schema)
        required_properties = []
        for x in range(0, length - 1):
            # parse a line in a schema
            tokens = lines_schema[x].split('"')
            if len(tokens) > 1:
                if tokens[1] == 'required':
                    for token in tokens:
                        if token == "]":
                            print ("correct end of required detected")
                        if token not in ignore_list:
                            if " " not in token:
                                required_properties.append(token)
        return required_properties

    def list_to_array(self, input_list):
        """
        generates an raml string representation of an python list
        :param input_list: python array
        :return: string as raml string representation. example = "[ 'blah', 'blah2' ]"
        """
        my_string = "["
        for x in input_list:
            comma = ", "
            my_string = my_string + '"' + x + '"' + comma
        # remove last comman (e.g. last 2 chars)
        my_string = my_string[:-2]
        my_string += "]"
        return my_string

    def list_to_string(self, input_list):
        """
        generates an raml string representation of an python list
        :rtype : string
        :param input_list: python array ["aa", "bb"
        :return: string example "aabb"
        """
        my_string = ""
        for x in input_list:
            my_string = my_string + x
        return my_string

    def list_resource(self, parse_tree, resource_name):
        """
        function to list the CRUDN behavior per resource
        e.g. it adds an entry to the CRUDN table
        :param parse_tree:
        :param resource_name: the resource_name

        """
        full_resource_name = "/" + str(resource_name)
        path = find_key_link(parse_tree, full_resource_name)

        row_cells = self.table.add_row().cells
        # row_cells[0].text = resource
        full_resource_name_no_query = full_resource_name.split("?")[0]
        row_cells[0].text = full_resource_name_no_query

        if path is not None:
            for method, mobj in path.items():
                # print "Method:",method
                # PUT == Create
                if method == "put":
                    row_cells[1].text = method
                # GET = Read
                if method == "get":
                    row_cells[2].text = method
                # POST - update  (agreed on 05/02/2015)
                if method == "post":
                    row_cells[3].text = method
                # DELETE = Delete
                if method == "delete":
                    row_cells[4].text = method
                # NOTIFY = NOTIFY (does not exist)
                if method == "get":
                    row_cells[5].text = "observe"

    def list_resources_crudn(self, parse_tree, resource_name=None):
        """
        function to create the CRUDN table
        :param parse_tree:
        :param resource_name:
        """
        level = 0
        
        # create the caption
        paragraph = self.document.add_paragraph('Table ', style='Caption')
        Table (paragraph)
        paragraph.add_run("The CRUDN operations of the resource with type 'rt' = "+resource_name)
        paragraph.style = 'TABLE-title'

        # create the table
        self.table = self.document.add_table(rows=1, cols=6)
        try:
            self.table.style = TABLESTYLE
        except:
            print ("no style set for table")
                
        hdr_cells = self.table.rows[0].cells
        hdr_cells[0].text = 'Resource'
        hdr_cells[1].text = 'Create'
        hdr_cells[2].text = 'Read'
        hdr_cells[3].text = 'Update'
        hdr_cells[4].text = 'Delete'
        hdr_cells[5].text = 'Notify'

        self.list_resource(parse_tree, resource_name)

    def list_properties(self, parse_tree, resource_name):
        """

        :param parse_tree:
        :param resource_name:
        """
        definitions = find_key_link(parse_tree, 'definitions')
        if definitions is not None:
            for object_name, json_object in definitions.items():
                print ("handling object:", object_name)
                property_list = find_key_link(json_object, 'properties')
                required_props = find_key_link(json_object, 'required')

                if property_list is not None:
                    for prop in property_list:
                        # fill the table
                        try:
                            if isinstance(property_list, dict):
                                if isinstance(property_list[prop], dict):
                                    print ("list_properties: property:", prop)
                                    description_text = property_list[prop].get('description', "")
                                    read_only = property_list[prop].get('readOnly')
                                    my_type = property_list[prop].get('type')
                                    if my_type is None:
                                        my_type = "multiple types: see schema"
                                    if my_type == "array":
                                        my_type += ": see schema"
                                    if my_type == "object":
                                        my_type += ": see schema"
                                    row_cells = self.tableAttribute.add_row().cells
                                    row_cells[0].text = str(prop)
                                    row_cells[1].text = str(my_type)
                                    if required_props is not None:
                                        if str(prop) in required_props:
                                            row_cells[2].text = "Yes"
                                        else:
                                            row_cells[2].text = "No"
                                    if read_only is not None and read_only is True:
                                        row_cells[3].text = "Read Only"
                                    if read_only is not None and read_only is False:
                                        row_cells[3].text = "Read Write"
                                    if read_only is None :
                                        # default value is false, see https://github.com/OAI/OpenAPI-Specification/blob/master/versions/2.0.md
                                        row_cells[3].text = "Read Write"
                                    text_par = row_cells[4].add_paragraph(description_text)
                                    text_par.alignment=WD_ALIGN_PARAGRAPH.LEFT
                                    #row_cells[4].text = description_text
                                else:
                                    print ("list_properties : not handled:", prop, property_list[prop])
                            else:
                                print ("list_properties : not handled:", prop, property_list[prop])

                        except:
                            traceback.print_exc()
                            pass


    def list_properties_derived(self, parse_tree, resource_name):
        """

        :param parse_tree:
        :param resource_name:
        """
        mydef = find_key_link(parse_tree, 'definitions')
        if mydef is not None:
            for object_name, json_object in mydef.items():
                print ("handling object:", object_name)
                if object_name == resource_name:
                    property_list = find_key_link(json_object, 'properties')
                    required_props = find_key_link(json_object, 'required')

                    if property_list is not None:
                        for prop in property_list:
                            # fill the table
                            try:
                                if isinstance(property_list, dict):
                                    if isinstance(property_list[prop], dict):
                                        print ("parse_schema: property:", prop)
                                        description_text = property_list[prop].get('description', "")
                                        ocf_resource = to_ocf = from_ocf = ""
                                        my_dict = property_list[prop].get("x-ocf-conversion")
                                        if my_dict is not None:
                                            ocf_resource = my_dict.get('x-ocf-alias', "")
                                            to_ocf = my_dict.get('x-to-ocf', "")
                                            from_ocf = my_dict.get('x-from-ocf', "")

                                        row_cells = self.tableAttribute.add_row().cells
                                        row_cells[0].text = str(prop)
                                        row_cells[1].text = str(ocf_resource)
                                        row_cells[2].text = self.list_to_string(to_ocf)
                                        row_cells[3].text = self.list_to_string(from_ocf)
                                        #row_cells[4].text = description_text
                                    else:
                                        print ("list_properties_derived : not handled:", prop, properties[prop])
                                else:
                                    print ("list_properties_derived : not handled:", prop, properties[prop])
                            except:
                                traceback.print_exc()
                                pass


    def list_properties_derived_table(self, parse_tree, resource_name):
        """

        :param parse_tree:
        :param resource_name:
        """
        mydef = find_key_link(parse_tree, 'definitions')
        if mydef is not None:
            for object_name, json_object in mydef.items():
                if object_name == resource_name:
                    print ("handling object:", object_name)
                    property_list = find_key_link(json_object, 'properties')
                    required_props = find_key_link(json_object, 'required')
                    if required_props is None:
                        required_props = find_key_link(parse_tree, 'required')
                    print ("list_properties_derived_table required properties:", required_props)
                    if required_props is None:
                        required_props = []

                    if property_list is not None:
                        for prop in property_list:
                            # fill the table
                            try:
                                if isinstance(property_list, dict):
                                    if isinstance(property_list[prop], dict):
                                        print ("parse_schema: property:", prop)
                                        description_text = property_list[prop].get('description', "")
                                        type_text = property_list[prop].get('type', "")
                                        row_cells = self.tableAttribute.add_row().cells
                                        row_cells[0].text = str(prop)
                                        row_cells[1].text = type_text
                                        if prop in required_props:
                                            row_cells[2].text = "yes"
                                        else:
                                            row_cells[2].text = "no"
                                        row_cells[3].text = description_text
                                    else:
                                        print ("list_properties_derived : not handled:", prop, properties[prop])
                                else:
                                    print ("list_properties_derived : not handled:", prop, properties[prop])
                            except:
                                traceback.print_exc()
                                pass
                            
                            
    def list_attributes(self, parse_tree, resource_name=None):
        """
        list all properties (attributes) in an table.
        create the table and fill it up
        :param parse_tree:
        :param resource_name:
        """
        
        # create the caption
        paragraph = self.document.add_paragraph('Table ', style='Caption')
        Table (paragraph)
        paragraph.add_run("The properties definitions of the resource with type 'rt' = "+resource_name)
        paragraph.style = 'TABLE-title'
        # create the table
        self.tableAttribute = self.document.add_table(rows=1, cols=5)
        try:
            self.tableAttribute.style = TABLESTYLE
        except:
            print ("no style set for table")

        #self.tableAttribute = self.document.add_table(rows=1, cols=5, style='TABLE-A')
        hdr_cells = self.tableAttribute.rows[0].cells
        hdr_cells[0].text = 'Property name'
        hdr_cells[1].text = 'Value type'
        hdr_cells[2].text = 'Mandatory'
        hdr_cells[3].text = 'Access mode'
        hdr_cells[4].text = 'Description'

        level = 1

        if resource_name is None:
            pass
        else:
            self.list_properties(parse_tree, resource_name )

        if self.sensor_switch is True:
            # auto generate the sensor value data..
            row_cells = self.tableAttribute.add_row().cells
            row_cells[0].text = "value"
            row_cells[1].text = "boolean"
            row_cells[2].text = "yes"
            row_cells[3].text = "Read Only"
            row_cells[4].text = "True = Sensed, False = Not Sensed."

    def list_attributes_derived(self, parse_tree, select_resource=None):

        """
        list all properties (attributes) in an table.
        create the table and fill it up
        :param parse_tree:
        :param select_resource:
        """
        
        # create the caption
        paragraph = self.document.add_paragraph('Table ', style='Caption')
        Table (paragraph)
        paragraph.add_run("The property mapping for "+select_resource+".")
        paragraph.style = 'TABLE-title'
        # create the table        
        #self.tableAttribute = self.document.add_table(rows=1, cols=4, style='TABLE-A')
        self.tableAttribute = self.document.add_table(rows=1, cols=4)
        try:
            self.tableAttribute.style = TABLESTYLE
        except:
            print ("no style set for table")
        hdr_cells = self.tableAttribute.rows[0].cells
        hdr_cells[0].text = str(self.derived_name) + ' Property name'
        hdr_cells[1].text = 'OCF Resource'
        hdr_cells[2].text = 'To OCF'
        hdr_cells[3].text = 'From OCF'
        level = 1
        if select_resource is None:
            pass
        else:
            self.list_properties_derived(parse_tree, select_resource )
            
        # create the caption
        paragraph = self.document.add_paragraph('Table ', style='Caption')
        Table (paragraph)
        paragraph.add_run("The properties of "+select_resource+".")
        paragraph.style = 'TABLE-title'
        # create the table        
        #self.tableAttribute = self.document.add_table(rows=1, cols=4, style='TABLE-A')
        self.tableAttribute = self.document.add_table(rows=1, cols=4)
        try:
            self.tableAttribute.style = TABLESTYLE
        except:
            print ("no style set for table")
            
        hdr_cells = self.tableAttribute.rows[0].cells
        hdr_cells[0].text = str(self.derived_name) + ' Property name'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Required'
        hdr_cells[3].text = 'Description'
        level = 1
        if select_resource is None:
            pass
        else:
            self.list_properties_derived_table(parse_tree, select_resource )

    def get_value_by_path_name(self, parse_tree, path_name, target):
        """
        retrieve the target key below the path_name
        :param parse_tree: tree to search from
        :param path_name: url name (without the /)
        :param target: key to find
        :return:
        """
        full_path_name = "/"+path_name
        json_path_dict = find_key_link(parse_tree, full_path_name)
        value = find_key_link(json_path_dict, target)
        return value

    def get_value_by_path_name2(self, parse_tree, path_name, target1, target2):
        """
        retrieve the target2 key below the target1 key below the path_name
        :param parse_tree: tree to search from
        :param path_name: url name (without the /)
        :param target1: key to find after path_name
        :param target1: key to find after target1 key
        :return:
        """
        full_path_name = "/"+path_name
        json_path_dict = find_key_link(parse_tree, full_path_name)
        value1 = find_key_link(json_path_dict, target1)
        value = find_key_link(value1, target2)
        return value


    def generate_sections(self, parse_tree, resource_name):
        """
        generate the individual sections
        :param parse_tree:
        :param resource_name:
        """
        title_name = find_key_link(parse_tree, 'title')
        print ("Title:", title_name)
        self.title = title_name
        par = self.document.add_heading(title_name, level=2)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading1'

        if resource_name is not None:
            rt_name = self.get_value_by_path_name(parse_tree, resource_name, "rt")
        else:
            # todo fix this
            rt_name = ""
        print ("RT   : `", rt_name, "`")

        # section Resource name (Introduction)
        if resource_name is not None:
            description_value = self.get_value_by_path_name2(parse_tree, resource_name, "get", "description")
        else:
            # TODO: fix this
            description_value = ""
        print ("description:", description_value)
        new_text = self.swag_unsanitize_description(description_value)
        print ("sanitized text:", new_text)
        if new_text != "":
            # add text + heading
            
            # section introduction
            par = self.document.add_heading('Introduction', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'
            par_text = self.document.add_paragraph(new_text)
            par_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
            #self.document.add_paragraph(value)

        # section URI
        if resource_name is not None:
            url_without_query= str(resource_name).split('?')[0]
        else:
            # TODO: fix this 
            url_without_query = ""
        if url_without_query != "":
            # write the data
            if self.annex_switch is False:
                par = self.document.add_heading('Example URI', level=3)
            else:
                par = self.document.add_heading('Wellknown URI', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'
            self.document.add_paragraph("/"+str(url_without_query))

        # section RT
        if resource_name is not None:
            par = self.document.add_heading('Resource type', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'
            if rt_name is not None:
                text = "The resource type (rt) is defined as: " + str(rt_name) + "."
                self.document.add_paragraph(text)
            else:
                print ("RT not found!")
        else:
            # derived model
            # list all definitions
            par = self.document.add_heading('Derived model', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'
            for def_name, def_data in parse_tree["definitions"].items():
                print ("derived model name (defintions):",def_name)
                if def_name is not None:
                    text = "The derived model: " + str(def_name) + ". "
                    description_text = def_data.get('description', "")
                    self.document.add_paragraph(text + description_text)
            par = self.document.add_heading('Property definition', level=3)
            for def_name, def_data in parse_tree["definitions"].items():
                print ("derived model name (table):",def_name)
                self.list_attributes_derived(parse_tree, select_resource=def_name)

        # section Swagger definition
        if resource_name is not None:
            par = self.document.add_heading('Swagger 2.0 definition', level=3)
        else:
            par = self.document.add_heading('Derived model definition', level=3)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading2'

        object_string = open(args.swagger, 'r').read()
        sanitized_text = self.swag_unsanitize_description(object_string)
        #object_string = json.dumps(parse_tree, sort_keys=True, indent=2, separators=(',', ': '))
        try:
            par = self.document.add_paragraph(sanitized_text, style='CODE-BLACK')
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except:
            pass

        if self.composite_switch is False:
            # do not add when the switch is true...
            # section property definition
            if resource_name is not None:
                par = self.document.add_heading('Property definition', level=3)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'
                self.list_attributes(parse_tree, resource_name=resource_name)

        if resource_name is not None:
            # section CRUDN definition
            par = self.document.add_heading('CRUDN behaviour', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'
            if resource_name is not None:
                self.list_resources_crudn(parse_tree, resource_name=resource_name)

            if self.schema_switch is True:
                # section extra JSON definition
                par = self.document.add_heading('Referenced JSON schemas', level=3)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'

                for my_schema_file in self.schema_files:
                    par = self.document.add_heading(my_schema_file, level=4)
                    if self.annex_switch is True:
                        par.style = 'ANNEX-heading2'
                    schema_text = open(my_schema_file, 'r').read()
                    try:
                        par = self.document.add_paragraph(self.add_justification("", schema_text), style='CODE-BLACK')
                        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    except:
                        pass

        if self.schemaWT_switch is True:
            # section extra JSON definition
            par = self.document.add_heading('Referenced JSON schemas', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'

            for schema_file in self.schemaWT_files:
                par = self.document.add_heading(schema_file, level=4)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'

                par = self.document.add_heading("Property definition", level=5)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'

                schema_text = open(schema_file, 'r').read()

                # create the caption
                paragraph = self.document.add_paragraph('Table ', style='Caption')
                Table (paragraph)
                paragraph.add_run("The properties definitions of schema file "+schema_file)
                paragraph.style = 'TABLE-title'
                
                # create the table
                #self.tableAttribute = self.document.add_table(rows=1, cols=5, style='TABLE-A')
                self.tableAttribute = self.document.add_table(rows=1, cols=5)
                try:
                    self.tableAttribute.style = TABLESTYLE
                except:
                    print ("no style set for table")

                hdr_cells = self.tableAttribute.rows[0].cells
                hdr_cells[0].text = 'Property name'
                hdr_cells[1].text = 'Value type'
                hdr_cells[2].text = 'Mandatory'
                hdr_cells[3].text = 'Access mode'
                hdr_cells[4].text = 'Description'

                # add fields in table with contents..
                self.parse_schema(schema_text)
                par = self.document.add_heading("Schema definition", level=5)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'
                try:
                    par = self.document.add_paragraph(self.add_justification("", schema_text), style='CODE-BLACK')
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except:
                    pass

    def convert(self):
        """
        conversion of the swagger data into the word document

        :return:
        """
        try:
            self.document = Document(docx=self.docx_name_in)
        except:
            print ("could not load file: ", self.docx_name_in)
            print ("make sure that docx file exist..")
            return

        self.generate_sections(self.json_parse_tree, self.resource_name)
        if self.docx_name_out is not None:
            self.document.save(self.docx_name_out)
            print ("document saved..", self.docx_name_out)


#
#   main of script
#
print ("************************")
print ("*** swagger2doc (v1.0.1) ***")
print ("************************")
parser = argparse.ArgumentParser()

parser.add_argument( "-ver"        , "--verbose"    , help="Execute in verbose mode", action='store_true')

parser.add_argument( "-swagger"    , "--swagger"    , default=None,
                     help="swagger file name",  nargs='?', const="", required=False)
parser.add_argument( "-schema"     , "--schema"     , default=None,
                     help="schema to be added to word document",  nargs='?', const="", required=False)
parser.add_argument( "-docx"       , "--docx"       , default=None,
                     help="word file in",  nargs='?', const="", required=False)
parser.add_argument( "-word_out"   , "--word_out"   , default=None,
                     help="word file out",  nargs='?', const="", required=False)
parser.add_argument( "-resource"   , "--resource"   , default=None,
                     help="resource (path) to be put in the word document",  nargs='?', const="", required=False)
parser.add_argument( "-schemadir"  , "--schemadir"  , default=".",
                     help="path to dir with additional referenced schemas",  nargs='?', const="", required=False)

parser.add_argument('-derived', '--derived', default=None, help='derived data model specificaton (--derived XXX) e.g. XXX Property Name in table use "." to ignore the property name setting')
parser.add_argument('-annex', '--annex', help='uses a annex heading instead of normal heading (--annex true)')

args = parser.parse_args()


print("file        : " + str(args.swagger))
print("resource    : " + str(args.resource))
print("schema      : " + str(args.schema))
print("schemadir   : " + str(args.schemadir))
print("docx        : " + str(args.docx))
print("word_out    : " + str(args.word_out))
print("derived     : " + str(args.derived))
print("")

try:
    #swagger_parser = SwaggerParser(swagger_path=args.swagger)  # Init with file
    swagger_parser=None
    worddoc = CreateWordDoc(swagger_parser)
    worddoc.docx_name_in = args.docx
    worddoc.docx_name_out = args.word_out
    worddoc.resource_name = args.resource

    annex_switch = args.annex
    if annex_switch is None:
        annex_switch = False
    else:
        annex_switch = True

    worddoc.annex_switch = annex_switch

    worddoc.derived_name = args.derived
    if worddoc.derived_name in ["."]:
        worddoc.derived_name = ""

    worddoc.convert()

    print (swagger_parser)
except:
    #print ("error in ", args.swagger)
    traceback.print_exc()
    #pass
